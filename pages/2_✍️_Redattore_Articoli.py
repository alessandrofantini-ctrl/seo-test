import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
import re
import json

st.set_page_config(page_title="SEO Copywriter - Connected Mode", layout="wide")

# =========================
# SIDEBAR (API KEY + SETTINGS)
# =========================
with st.sidebar:
    st.title("Configurazione")

    openai_api_key = st.text_input("OpenAI key", type="password")
    if not openai_api_key and "OPENAI_API_KEY" in st.secrets:
        openai_api_key = st.secrets["OPENAI_API_KEY"]

    st.markdown("---")
    brand_name = st.text_input("Nome azienda/brand (opzionale)", placeholder="Es. Nome azienda")
    target_page_url = st.text_input("URL pagina/servizio per CTA (opzionale)", placeholder="https://...")

    st.markdown("---")
    lunghezza_label = st.select_slider(
        "Lunghezza articolo",
        options=["Standard", "Long form", "Authority guide (pillar)"],
        value="Long form"
    )

    creativita = st.slider("Creatività", 0.0, 1.0, 0.35)

    st.markdown("---")
    st.subheader("Keyword secondarie (fallback)")
    secondary_kw_manual = st.text_area(
        "Se il brief non contiene il JSON o è incompleto, inserisci qui (una per riga)",
        height=120
    )

    st.markdown("---")
    st.caption("Maiuscole: lo script forza sentence case per titoli (no title case).")

# =========================
# IMPORT AUTOMATICO BRIEF DA ANALISI SEO
# =========================
brief_default = ""
json_data = None

if "ultimo_brief" in st.session_state:
    st.success("Brief importato automaticamente dal tool Analisi SEO.")
    brief_default = st.session_state["ultimo_brief"]

    match = re.search(r"```json\s*(\{.*?\})\s*```", brief_default, re.S)
    if match:
        try:
            json_data = json.loads(match.group(1))
        except Exception:
            json_data = None

st.title("SEO copywriter – connected mode")

brief_input = st.text_area(
    "Brief SEO (auto-importato, modificabile)",
    value=brief_default,
    height=420
)

# =========================
# LUNGHEZZA TARGET
# =========================
def get_word_target(label):
    if label == "Standard":
        return "1200-1600"
    elif label == "Long form":
        return "1800-2500"
    else:
        return "2500-3500"

# =========================
# DOCX EXPORT
# =========================
def create_docx(content):
    doc = Document()
    doc.add_heading("SEO article", 0)
    for line in content.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def parse_secondary_keywords_from_json(j):
    if not j:
        return []
    secs = j.get("secondary_keywords", [])
    if isinstance(secs, list):
        return [str(x).strip() for x in secs if str(x).strip()]
    return []

def parse_outline_from_json(j):
    """
    Atteso:
    "outline": [{"h2":"...", "h3":["...","..."]}, ...]
    """
    if not j:
        return None
    outline = j.get("outline")
    if isinstance(outline, list) and outline:
        ok = True
        for item in outline:
            if not isinstance(item, dict) or "h2" not in item:
                ok = False
                break
        return outline if ok else None
    return None

# =========================
# GENERAZIONE ARTICOLO
# =========================
if st.button("Genera articolo SEO"):
    if not openai_api_key or not brief_input:
        st.error("Inserisci OpenAI key e brief.")
        st.stop()

    client = OpenAI(api_key=openai_api_key)
    word_target = get_word_target(lunghezza_label)

    # Strategia lingua/keyword dal JSON (se presente)
    primary_kw = ""
    language = "it"
    outline_json = None
    secondary_kws = []

    if json_data:
        primary_kw = json_data.get("primary_keyword", "") or json_data.get("primary", "") or ""
        language = (json_data.get("language", "") or "it").lower()
        outline_json = parse_outline_from_json(json_data)
        secondary_kws = parse_secondary_keywords_from_json(json_data)

    # fallback keyword secondarie manuali
    if secondary_kw_manual.strip():
        manual = [x.strip() for x in secondary_kw_manual.splitlines() if x.strip()]
        # non duplicare
        for k in manual:
            if k not in secondary_kws:
                secondary_kws.append(k)
    secondary_kws = secondary_kws[:20]

    # brand fallback
    brand = (brand_name or "").strip()

    # prompt: riduce “fuffa”, forza decision making e sentence case
    system_prompt = f"""
Sei un senior SEO copywriter.

Scrivi contenuti autorevoli ma concreti, senza frasi generiche.
Stile: chiaro, operativo, orientato a decisioni e casi reali.

Regole:
- Scrivi in italiano se possibile (language={language}). Se language è en/es/fr/de, scrivi in quella lingua.
- Maiuscole: usa sentence case per H1/H2/H3 e frasi. Non usare title case.
- Evita claim numerici non supportati (no % casuali). Se dai numeri, dichiarali come "stima indicativa" o "esempio".
- Mantieni struttura Markdown pulita.
- Niente “come vedremo”, “nel mondo di oggi”, “rivoluzionare” ripetuto.
"""

    # Se l’outline è nel JSON, imponiamo di seguirlo 1:1 (molto più affidabile)
    outline_block = ""
    if outline_json:
        outline_block = json.dumps(outline_json, ensure_ascii=False, indent=2)
    else:
        outline_block = "Non disponibile: segui l'outline scritto nel brief."

    user_prompt = f"""
Brief SEO:
{brief_input}

Dati fissi:
- primary keyword: "{primary_kw}"
- brand: "{brand}" (se vuoto, non inserirlo forzatamente)
- keyword secondarie (usa solo se coerenti): {secondary_kws if secondary_kws else "nessuna"}
- lunghezza minima: {word_target} parole
- url CTA (se presente): "{target_page_url}"

Obiettivo:
Scrivi l'articolo completo, pratico e non generico.

Requisiti obbligatori:
1) Introduzione 150-220 parole:
   - include la primary keyword una sola volta in modo naturale
   - promette cosa impara il lettore e per chi è utile
   - evita frasi vaghe

2) Struttura:
   - Rispetta H1/H2/H3 del brief.
   - Se trovi un outline nel JSON, seguilo esattamente.
   Outline JSON (se presente):
{outline_block}

3) Per ogni H2:
   - includi almeno: esempio concreto, checklist o mini-framework (3-7 bullet), e un “errore comune da evitare”.
   - se l'intento è informativo, mantieni tono pratico.
   - se l'intento è commerciale, inserisci criteri di scelta e segnali di ROI.

4) Tabelle:
   - includi massimo 1-2 tabelle, solo se aiutano una decisione (non tabelle enciclopediche).

5) FAQ:
   - 5 domande utili, orientate a dubbi reali (costi, tempi, dati, rischi, governance), non definizioni scolastiche.

6) CTA finale:
   - 2-3 CTA brevi e concrete.
   - Se c'è un URL CTA, invoglia a visitarlo senza spam.

Output:
- Solo Markdown dell'articolo, senza commenti extra.
"""

    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=creativita,
        max_tokens=7500
    )

    article = resp.choices[0].message.content

    st.markdown("## Articolo generato")
    st.markdown(article)

    docx = create_docx(article)
    st.download_button(
        "Scarica articolo (.docx)",
        data=docx,
        file_name="seo_article.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
