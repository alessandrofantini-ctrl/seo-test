import streamlit as st
import requests
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from io import BytesIO
import json
import re
from urllib.parse import urlparse
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import Counter
from requests.adapters import HTTPAdapter, Retry

# --- CONFIGURAZIONE ---
st.set_page_config(page_title="SEO Brief Generator", layout="wide")

# --- DIZIONARIO MERCATI ---
MARKETS = {
    "🇮🇹 Italia": {"gl": "it", "hl": "it", "domain": "google.it"},
    "🇺🇸 USA (English)": {"gl": "us", "hl": "en", "domain": "google.com"},
    "🇬🇧 UK": {"gl": "uk", "hl": "en", "domain": "google.co.uk"},
    "🇪🇸 Spagna": {"gl": "es", "hl": "es", "domain": "google.es"},
    "🇫🇷 Francia": {"gl": "fr", "hl": "fr", "domain": "google.fr"},
    "🇩🇪 Germania": {"gl": "de", "hl": "de", "domain": "google.de"},
}

# =========================
# HTTP SESSION + RETRY
# =========================
def build_session():
    s = requests.Session()
    retries = Retry(
        total=3,
        backoff_factor=0.6,
        status_forcelist=(429, 500, 502, 503, 504),
        allowed_methods=("GET",),
        raise_on_status=False,
    )
    s.mount("https://", HTTPAdapter(max_retries=retries))
    s.mount("http://", HTTPAdapter(max_retries=retries))
    return s

HTTP = build_session()

UA = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
}

# =========================
# UTILS
# =========================
def safe_text(s: str) -> str:
    if not s:
        return ""
    return re.sub(r"\s+", " ", s).strip()

def truncate_chars(s: str, n: int) -> str:
    s = safe_text(s)
    return s[:n].rstrip()

def domain_of(url: str) -> str:
    try:
        return urlparse(url).netloc.replace("www.", "")
    except Exception:
        return ""

def extract_json_ld(soup: BeautifulSoup):
    scripts = soup.find_all("script", type="application/ld+json")
    out = []
    for sc in scripts[:12]:
        txt = sc.get_text(strip=True)
        if not txt:
            continue
        try:
            out.append(json.loads(txt))
        except Exception:
            pass
    return out

def remove_boilerplate(soup: BeautifulSoup):
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe"]):
        tag.decompose()
    for selector in ["nav", "header", "footer", "aside", "form"]:
        for tag in soup.select(selector):
            tag.decompose()
    for cls in ["cookie", "cookies", "cookie-banner", "newsletter", "modal", "popup"]:
        for tag in soup.select(f".{cls}"):
            tag.decompose()
    return soup

def detect_main_container(soup: BeautifulSoup):
    for tag in ["article", "main"]:
        el = soup.find(tag)
        if el:
            txt = el.get_text(" ", strip=True)
            if txt and len(txt) > 600:
                return el
    return soup.body if soup.body else soup

def normalize_sentence_case(text: str) -> str:
    t = safe_text(text)
    if not t:
        return ""
    if t.isupper() and len(t) > 6:
        t = t.lower()
    return t[0].upper() + t[1:] if len(t) > 1 else t.upper()

STOPWORDS = set("""
a al allo alla alle agli all' and are as at be by con che da dal dalla dalle degli dei del della delle di do
e ed en est et for from il in is it la le lo los las les more nel nei nell' of on or per por pour que qui
su the to un una uno und une with y zu""".split())

def tokenize(text: str):
    text = (text or "").lower()
    tokens = re.findall(r"[a-zàèéìòùäöüßñç0-9]{3,}", text, flags=re.I)
    return [t for t in tokens if t not in STOPWORDS]

# =========================
# CACHING
# =========================
@st.cache_data(show_spinner=False, ttl=60 * 60)
def get_serp_data(query, api_key, gl, hl, domain):
    params = {
        "engine": "google",
        "q": query,
        "api_key": api_key,
        "hl": hl,
        "gl": gl,
        "google_domain": domain
    }
    r = HTTP.get("https://serpapi.com/search", params=params, timeout=25)
    if r.status_code != 200:
        return None
    try:
        return r.json()
    except Exception:
        return None

@st.cache_data(show_spinner=False, ttl=60 * 60)
def scrape_site_content(url, include_meta=True, include_schema=True, max_text_chars=9000):
    """
    Estrae segnali SEO + contenuto competitor (testo ripulito, headings, liste).
    """
    data = {
        "url": url,
        "domain": domain_of(url),
        "title": "",
        "meta_description": "",
        "canonical": "",
        "h1": "",
        "h2": [],
        "h3": [],
        "word_count": 0,
        "text_sample": "",
        "top_terms": [],
        "lang": "",
        "schema_types": [],
        "has_faq_schema": False,
        "question_headings": [],
    }

    try:
        resp = HTTP.get(url, headers=UA, timeout=18, allow_redirects=True)
        if resp.status_code >= 400 or not resp.text:
            return None

        soup = BeautifulSoup(resp.text, "html.parser")
        soup = remove_boilerplate(soup)

        html_tag = soup.find("html")
        if html_tag and html_tag.get("lang"):
            data["lang"] = safe_text(html_tag.get("lang"))

        if include_meta:
            if soup.title and soup.title.string:
                data["title"] = safe_text(soup.title.string)
            md = soup.find("meta", attrs={"name": "description"})
            if md and md.get("content"):
                data["meta_description"] = safe_text(md.get("content"))
            canon = soup.find("link", rel=lambda x: x and "canonical" in x.lower())
            if canon and canon.get("href"):
                data["canonical"] = safe_text(canon.get("href"))

        main = detect_main_container(soup)

        # headings
        for tag in main.find_all(["h1", "h2", "h3"])[:80]:
            txt = safe_text(tag.get_text(" ", strip=True))
            if not txt:
                continue
            if tag.name == "h1" and not data["h1"]:
                data["h1"] = txt
            elif tag.name == "h2" and len(data["h2"]) < 30:
                data["h2"].append(txt)
            elif tag.name == "h3" and len(data["h3"]) < 45:
                data["h3"].append(txt)

            if "?" in txt:
                data["question_headings"].append(txt)

        # testo: paragrafi + liste
        paragraphs = main.find_all("p")
        lis = main.find_all("li")

        p_text = " ".join([safe_text(p.get_text(" ", strip=True)) for p in paragraphs])
        li_text = " ".join([safe_text(li.get_text(" ", strip=True)) for li in lis[:140]])

        text_content = safe_text((p_text + " " + li_text).strip())
        if len(text_content) > max_text_chars:
            text_content = text_content[:max_text_chars]

        data["word_count"] = len(text_content.split()) if text_content else 0
        data["text_sample"] = text_content[:2400]

        # top terms
        toks = tokenize(text_content)
        common = Counter(toks).most_common(25)
        data["top_terms"] = [t for t, _ in common]

        # schema
        if include_schema:
            jlds = extract_json_ld(soup)
            types = set()
            has_faq = False
            for item in jlds:
                items = item if isinstance(item, list) else [item]
                for it in items:
                    if isinstance(it, dict) and "@type" in it:
                        t = it["@type"]
                        if isinstance(t, list):
                            for tt in t:
                                types.add(str(tt))
                        else:
                            types.add(str(t))
                        if "FAQPage" in str(t):
                            has_faq = True
            data["schema_types"] = sorted(types)[:25]
            data["has_faq_schema"] = has_faq

        return data
    except Exception:
        return None

def build_serp_snapshot(serp_json, max_items):
    snapshot = {
        "organic": [],
        "paa": [],
        "related_searches": [],
        "features": [],
    }
    if not serp_json:
        return snapshot

    for res in serp_json.get("organic_results", [])[:max_items]:
        snapshot["organic"].append({
            "position": res.get("position"),
            "title": res.get("title"),
            "link": res.get("link"),
            "snippet": res.get("snippet") or res.get("snippet_highlighted_words"),
            "source": res.get("source"),
        })

    for q in serp_json.get("related_questions", [])[:20]:
        if q.get("question"):
            snapshot["paa"].append(q.get("question"))

    for r in serp_json.get("related_searches", [])[:20]:
        if r.get("query"):
            snapshot["related_searches"].append(r.get("query"))

    if serp_json.get("answer_box"):
        snapshot["features"].append("answer_box")
    if serp_json.get("knowledge_graph"):
        snapshot["features"].append("knowledge_graph")
    if serp_json.get("shopping_results"):
        snapshot["features"].append("shopping_results")
    if serp_json.get("local_results"):
        snapshot["features"].append("local_results")
    if serp_json.get("top_stories"):
        snapshot["features"].append("top_stories")
    if serp_json.get("inline_videos"):
        snapshot["features"].append("inline_videos")

    return snapshot

def create_docx_from_markdownish(content_md: str, kw: str):
    doc = Document()
    doc.add_heading(f"SEO brief: {kw}", 0)
    for line in content_md.splitlines():
        l = line.strip()
        if not l:
            doc.add_paragraph("")
            continue
        if l.startswith("### "):
            doc.add_heading(l.replace("### ", ""), level=3)
        elif l.startswith("## "):
            doc.add_heading(l.replace("## ", ""), level=2)
        elif l.startswith("# "):
            doc.add_heading(l.replace("# ", ""), level=1)
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def aggregate_competitor_insights(competitors, target_lang):
    h2_all, terms_all, q_all = [], [], []

    for c in competitors:
        for h2 in c.get("h2", [])[:30]:
            h = safe_text(h2)
            if h:
                h2_all.append(h.lower())
        for t in c.get("top_terms", [])[:25]:
            terms_all.append(t.lower())
        for q in c.get("question_headings", [])[:25]:
            qq = safe_text(q)
            if qq:
                q_all.append(qq.lower())

    def norm_heading(h):
        h = re.sub(r"\s+", " ", h)
        h = re.sub(r"[^\w\sàèéìòùäöüßñç-]", "", h)
        return h.strip()

    top_h2 = [normalize_sentence_case(x) for x, _ in Counter([norm_heading(x) for x in h2_all if x]).most_common(12)]
    top_terms = [x for x, _ in Counter([x for x in terms_all if x and x not in STOPWORDS]).most_common(20)]
    top_q = [normalize_sentence_case(x) for x, _ in Counter([norm_heading(x) for x in q_all if x]).most_common(8)]

    return {"top_h2": top_h2, "top_terms": top_terms, "top_questions": top_q}

# =========================
# SIDEBAR
# =========================
with st.sidebar:
    st.title("⚙️ SEO settings")

    # BLOCCO CHIAVI (come richiesto)
    openai_api_key = st.text_input("OpenAI key", type="password")
    serp_api_key = st.text_input("SerpApi key", type="password")
    if not openai_api_key and "OPENAI_API_KEY" in st.secrets:
        openai_api_key = st.secrets["OPENAI_API_KEY"]
    if not serp_api_key and "SERP_API_KEY" in st.secrets:
        serp_api_key = st.secrets["SERP_API_KEY"]

    st.markdown("---")
    st.subheader("🌍 Mercato")
    selected_market_label = st.selectbox("Seleziona mercato target", list(MARKETS.keys()))
    market_params = MARKETS[selected_market_label]

    st.markdown("---")
    st.subheader("🏷️ Brand")
    brand_name = st.text_input("Nome azienda/brand (per meta title/description)", placeholder="Es. Nome azienda")

    st.markdown("---")
    st.subheader("🎯 Target cliente")
    client_url = st.text_input("URL sito cliente (opzionale)", placeholder="https://www.tuosito.it")
    custom_usp = st.text_area("USP / punti di forza (opzionale)", height=90)
    tone_of_voice = st.selectbox("Tono di voce", ["Autorevole & tecnico", "Empatico & problem solving", "Diretto & commerciale"])

    st.markdown("---")
    st.subheader("🔑 Keyword secondarie (opzionale)")
    secondary_keywords_manual = st.text_area("Una per riga", height=120)

    st.markdown("---")
    st.subheader("🔧 Analisi")
    deep_mode = st.toggle("Deep mode (più competitor)", value=True)
    max_competitors = st.slider("Competitor da analizzare", 2, 10, 6 if deep_mode else 3)
    max_workers = st.slider("Parallelismo scraping", 2, 10, 6)
    include_schema = st.toggle("Estrai schema (JSON-LD)", value=True)
    include_meta = st.toggle("Estrai meta (title/description/canonical)", value=True)

    st.markdown("---")
    st.subheader("🤖 Modello")
    model_name = st.selectbox("Modello", ["gpt-4o", "gpt-4o-mini"], index=0)
    include_json_block = st.toggle("Includi blocco JSON finale", value=False)

# =========================
# MAIN
# =========================
st.title("🚀 SEO brief generator multi-country")
st.markdown(f"Analisi impostata su: **{selected_market_label}**")

col1, col2 = st.columns([3, 1])
with col1:
    keyword = st.text_input("Keyword principale", placeholder="Es. impianti elettrici industriali")
with col2:
    target_intent = st.selectbox("Intento", ["Informativo", "Commerciale", "Navigazionale"])

st.caption("Nota: SerpApi fornisce la SERP (title/snippet/link). Il contenuto dei competitor viene letto via scraping HTML.")

# =========================
# ACTION
# =========================
if st.button("Avvia analisi"):
    if not keyword or not openai_api_key or not serp_api_key:
        st.error("Inserisci keyword e API keys.")
        st.stop()

    status = st.status(f"Avvio scansione su {selected_market_label}…", expanded=True)

    # 1) SERP
    status.write("🔍 Analisi SERP…")
    serp = get_serp_data(
        keyword,
        serp_api_key,
        gl=market_params["gl"],
        hl=market_params["hl"],
        domain=market_params["domain"]
    )
    if not serp or "organic_results" not in serp:
        status.update(label="Errore SerpApi", state="error")
        st.error("Nessun dato SERP. Verifica SerpApi key o query.")
        st.stop()

    serp_snapshot = build_serp_snapshot(serp, max_competitors)
    organic_urls = [x["link"] for x in serp_snapshot["organic"] if x.get("link")][:max_competitors]

    with st.expander("SERP snapshot (debug)", expanded=False):
        st.json(serp_snapshot, expanded=2)

    # 2) CLIENT (opzionale, light)
    client_context = ""
    if client_url:
        status.write("🏢 Lettura sito cliente (light)…")
        cd = scrape_site_content(client_url, include_meta=True, include_schema=False, max_text_chars=4000)
        if cd:
            client_context = (
                f"Brand site title: {truncate_chars(cd.get('title',''), 120)}\n"
                f"Brand meta description: {truncate_chars(cd.get('meta_description',''), 200)}\n"
                f"Brand H1: {truncate_chars(cd.get('h1',''), 120)}\n"
                f"Brand excerpt: {truncate_chars(cd.get('text_sample',''), 520)}"
            )
        else:
            client_context = "Sito cliente non leggibile o bloccato."
    else:
        client_context = "Nessun sito cliente fornito."

    if custom_usp:
        client_context += f"\nUSP: {truncate_chars(custom_usp, 450)}"

    # 3) SCRAPING COMPETITOR (contenuto reale)
    status.write("⚔️ Scraping competitor (contenuto reale)…")
    competitor_results = []
    prog = status.progress(0.0)

    futures = []
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        for u in organic_urls:
            futures.append(ex.submit(scrape_site_content, u, include_meta, include_schema))

        done = 0
        for fut in as_completed(futures):
            done += 1
            prog.progress(done / max(1, len(futures)))
            comp = fut.result()
            if comp:
                competitor_results.append(comp)

    prog.empty()

    if not competitor_results:
        status.update(label="Errore scraping", state="error")
        st.error("Non sono riuscito a leggere i competitor (403/timeout). Prova a ridurre competitor o parallelismo.")
        st.stop()

    # 🔎 DEBUG COMPETITOR (come SERP snapshot)
    with st.expander("Competitor snapshot (debug)", expanded=False):
        # ordina per presenza contenuto/word_count
        competitor_results_sorted = sorted(competitor_results, key=lambda x: x.get("word_count", 0), reverse=True)
        for c in competitor_results_sorted:
            label = f"{c.get('domain','') or 'competitor'} — {truncate_chars(c.get('title',''), 70) or c.get('url','')}"
            with st.expander(label, expanded=False):
                debug_obj = {
                    "url": c.get("url"),
                    "domain": c.get("domain"),
                    "lang": c.get("lang"),
                    "meta": {
                        "title": c.get("title"),
                        "meta_description": c.get("meta_description"),
                        "canonical": c.get("canonical"),
                    },
                    "headings": {
                        "h1": c.get("h1"),
                        "h2": c.get("h2", [])[:20],
                        "h3": c.get("h3", [])[:25],
                        "question_headings": c.get("question_headings", [])[:15],
                    },
                    "content": {
                        "word_count": c.get("word_count", 0),
                        "top_terms": c.get("top_terms", [])[:25],
                        "text_sample": c.get("text_sample", ""),
                    },
                    "schema": {
                        "schema_types": c.get("schema_types", [])[:25],
                        "has_faq_schema": c.get("has_faq_schema", False),
                    },
                }
                st.json(debug_obj, expanded=2)

    # 4) AGGREGA INSIGHTS
    status.write("🧩 Aggregazione insight competitor…")
    agg = aggregate_competitor_insights(competitor_results, market_params["hl"])

    # 5) Keyword secondarie manuali
    manual_secs = []
    if secondary_keywords_manual:
        manual_secs = [safe_text(x) for x in secondary_keywords_manual.splitlines() if safe_text(x)]
        manual_secs = manual_secs[:25]

    # 6) Prompt: outline H2/H3 obbligatorio
    status.write("🧠 Generazione brief…")

    brand = safe_text(brand_name) if brand_name else ""
    target_lang = market_params["hl"]

    competitor_compact = []
    for c in competitor_results[:max_competitors]:
        competitor_compact.append({
            "url": c.get("url"),
            "title": truncate_chars(c.get("title", ""), 120),
            "h1": truncate_chars(c.get("h1", ""), 120),
            "h2": [truncate_chars(x, 90) for x in (c.get("h2") or [])[:12]],
            "h3": [truncate_chars(x, 90) for x in (c.get("h3") or [])[:12]],
            "word_count": c.get("word_count", 0),
        })

    system_prompt = (
        "Sei un Senior SEO strategist. Produci brief pratici e brevi, orientati all'esecuzione. "
        "Non inserire teoria: solo ciò che serve per scrivere una pagina migliore dei competitor."
    )

    user_prompt = f"""
Keyword principale: "{keyword}"
Mercato: {selected_market_label}
Lingua output SEO (meta/h1/h2/h3): {target_lang}
Intento: {target_intent}
Tono di voce: {tone_of_voice}
Brand name: "{brand}" (se vuoto, non forzare il brand nel title)

Keyword secondarie già fornite (se presenti): {manual_secs if manual_secs else "Nessuna"}

SERP:
- Features: {", ".join(serp_snapshot.get("features", [])) if serp_snapshot.get("features") else "N/D"}
- PAA: {serp_snapshot.get("paa", [])[:10] if serp_snapshot.get("paa") else []}
- Related searches: {serp_snapshot.get("related_searches", [])[:12] if serp_snapshot.get("related_searches") else []}

Competitor (sintesi):
{json.dumps(competitor_compact, ensure_ascii=False)}

Pattern competitor:
- H2 ricorrenti: {agg.get("top_h2", [])}
- Termini ricorrenti: {agg.get("top_terms", [])[:18]}
- Domande ricorrenti: {agg.get("top_questions", [])}

Contesto cliente:
{client_context}

REGOLE IMPORTANTI:
1) Output in ITALIANO per le istruzioni, ma meta title/description/H1/H2/H3 nella lingua target: {target_lang}.
2) Maiuscole: sentence case per title/description/H1/H2/H3 (solo prima lettera maiuscola).
   Non mettere la maiuscola a ogni parola. Mantieni acronimi e brand corretti.
3) Meta title:
   - preferisci "keyword | Brand" se Brand presente
   - max 60 caratteri
   - 3 varianti (v1/v2/v3)
4) Meta description:
   - max 155 caratteri
   - 3 varianti (v1 informativa, v2 conversione, v3 ibrida)
5) Output compatto: niente spiegoni.

FORMATTO DI RISPOSTA (obbligatorio):
## meta
- title (v1): ...
- title (v2): ...
- title (v3): ...
- description (v1): ...
- description (v2): ...
- description (v3): ...

## h1
- ...

## outline (H2/H3)
Scrivi massimo 10 H2. Per ogni H2 scrivi 2–4 H3.
Formato:
- H2: ...
  - H3: ...
  - H3: ...
  - Nota (IT): una riga su cosa inserire e quali prove/esempi usare.

Integra PAA e correlate dove utile.

## keyword set
- primary: ...
- secondary (max 12): ... (usa anche quelle manuali se coerenti)

## faq
5 domande (in {target_lang}) + risposta 1 frase (in {target_lang})

## cta
3 CTA brevi (in italiano) coerenti con intento "{target_intent}"

{"## json\nIncludi un blocco JSON minimale (solo meta, h1, outline, secondary, faq) in ```json```." if include_json_block else ""}
"""

    client = OpenAI(api_key=openai_api_key)
    resp = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.5
    )
    output = resp.choices[0].message.content

    status.update(label="Brief pronto!", state="complete", expanded=False)

    st.markdown(output)

    docx = create_docx_from_markdownish(output, keyword)
    st.download_button(
        "📥 Scarica brief .docx",
        docx,
        f"brief_{keyword.replace(' ','_')}.docx"
    )

    json_match = re.search(r"```json\s*(\{.*?\})\s*```", output, flags=re.S)
    if json_match:
        try:
            parsed = json.loads(json_match.group(1))
            st.download_button(
                "📥 Scarica brief JSON",
                data=json.dumps(parsed, ensure_ascii=False, indent=2).encode("utf-8"),
                file_name=f"brief_{keyword.replace(' ','_')}.json",
                mime="application/json"
            )
        except Exception:
            pass
